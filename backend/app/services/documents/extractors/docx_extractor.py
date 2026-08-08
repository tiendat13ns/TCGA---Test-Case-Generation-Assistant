from app.services.documents.extractors.base import BaseExtractor

# Mapping từ Word Heading style → Markdown prefix
_HEADING_STYLE_TO_MARKDOWN = {
    "heading 1": "#",
    "heading 2": "##",
    "heading 3": "###",
    "heading 4": "####",
}

# Nhiều tài liệu SRS/Use Case thực tế KHÔNG dùng Word Heading style mà chỉ bôi đậm
# (bold) tên section thủ công (vd: "Mục tiêu:", "Luồng chính:", "Luồng ngoại lệ:") —
# đã xác nhận trên tài liệu thật của dự án. Heading kiểu này được coi là cấp 3.
_BOLD_HEADING_MARKDOWN = "###"
_BOLD_LABEL_MAX_CHARS = 60


def _split_bold_label(paragraph) -> tuple[str, str] | None:
    """Nếu đoạn văn bắt đầu bằng một cụm text ngắn được bôi đậm (label) rồi mới tới
    nội dung thường (thường ngăn nhau bởi xuống dòng mềm Shift+Enter trong CÙNG một
    paragraph), tách heading ra khỏi phần thân. Trả về (label, phần_thân_còn_lại)
    hoặc None nếu đoạn văn không khớp pattern này.
    """
    runs = [r for r in paragraph.runs if r.text]
    if not runs:
        return None

    first = runs[0]
    label = first.text.strip()
    if not label or len(label) > _BOLD_LABEL_MAX_CHARS or not first.bold:
        return None

    rest_runs = runs[1:]
    if not rest_runs:
        # Cả đoạn chỉ có 1 run và toàn bộ đều bôi đậm — cả đoạn chính là heading.
        return (label, "")

    # Bỏ qua các run rỗng/chỉ chứa khoảng trắng khi kiểm tra bold — Word đôi khi tách
    # riêng ký tự xuống dòng thành 1 run bôi đậm dù nội dung thân bài phía sau không
    # hề bôi đậm.
    non_ws_rest = [r for r in rest_runs if r.text.strip()]
    if any(r.bold for r in non_ws_rest):
        # Có nội dung bôi đậm khác ngoài label — nhiều khả năng là 1 câu bôi đậm dài
        # bình thường, không phải label + body.
        return None

    rest_text = "".join(r.text for r in rest_runs)
    if not rest_text.startswith("\n"):
        # Label không đứng riêng 1 dòng — không đủ tin cậy để coi là heading.
        return None

    body = rest_text.lstrip("\n").strip()
    return (label, body)


# Fallback cho các label KHÔNG được bôi đậm (nên _split_bold_label bỏ sót) nhưng vẫn
# là tên section SRS/Use Case phổ biến, đứng riêng ở đầu dòng dạng "Tên section:".
# Chỉ khớp khi CẢ dòng đầu tiên của đoạn văn == đúng 1 label trong danh sách này —
# so khớp CHÍNH XÁC (không phải substring) để tránh nhầm với câu bình thường kết
# thúc bằng dấu ":" (vd: "Người dùng nhập thông tin project gồm:").
_KNOWN_SECTION_LABELS = {
    # English
    "input", "inputs", "output", "outputs",
    "precondition", "preconditions", "pre-condition", "pre-conditions",
    "postcondition", "postconditions", "post-condition", "post-conditions",
    "acceptance criteria", "business rules", "business rule",
    "priority", "actor", "actors", "trigger", "triggers",
    "main flow", "alternative flow", "alternative flows",
    "exception flow", "exception flows", "assumptions", "constraints",
    "dependencies", "notes", "note", "description", "objective", "objectives",
    "scope", "frequency of use", "data to be stored",
    # Vietnamese
    "mục tiêu", "mô tả", "actor chính", "actor phụ", "luồng chính",
    "luồng thay thế", "luồng ngoại lệ", "đầu vào", "đầu ra",
    "dữ liệu cần lưu", "tần suất sử dụng", "độ ưu tiên", "ưu tiên",
    "tiêu chí chấp nhận", "ghi chú", "giả định", "ràng buộc",
}


def _match_known_section_label(text: str) -> tuple[str, str] | None:
    """Nhận diện heading dựa trên danh sách tên section SRS/Use Case phổ biến. Chỉ
    xét dòng đầu tiên (trước xuống dòng đầu tiên nếu có) — bao quát luôn trường hợp
    label + body cùng nằm trong 1 paragraph, ngăn nhau bởi Shift+Enter."""
    first_line, _, rest = text.partition("\n")
    first_line = first_line.strip()
    if not first_line.endswith(":"):
        return None
    label = first_line[:-1].strip()
    if label.lower() in _KNOWN_SECTION_LABELS:
        return (first_line, rest.strip())
    return None


class DocxExtractor(BaseExtractor):
    """Trích xuất text từ file .docx, giữ nguyên cấu trúc Heading dưới dạng Markdown.

    Chiến lược:
    - Các đoạn văn mang Word style Heading 1/2/3/4 được chuyển thành
      thẻ Markdown # / ## / ### / #### để MarkdownHeaderTextSplitter có thể
      nhận diện ranh giới chunk sau này.
    - Nếu không có Heading style, thử suy luận heading từ đoạn text ngắn được bôi đậm
      đứng đầu dòng (xem _split_bold_label) — bù cho các tài liệu chỉ bôi đậm thủ công.
    - Nếu vẫn không khớp (label không bôi đậm), thử khớp với danh sách tên section
      SRS/Use Case phổ biến (xem _match_known_section_label), ví dụ "Input:",
      "Output:", "Acceptance Criteria:".
    - Nội dung bảng (table) vẫn được giữ nguyên dạng pipe-delimited.
    """

    def extract(self, file_path: str) -> str:
        try:
            from docx import Document
        except ImportError as exc:
            raise RuntimeError("python-docx is required to extract DOCX files") from exc

        try:
            document = Document(file_path)
        except Exception as exc:
            raise ValueError(f"Could not read DOCX file: {exc}") from exc

        parts: list[str] = []

        for paragraph in document.paragraphs:
            text = paragraph.text.strip()
            if not text:
                continue

            style_name = ""
            if paragraph.style:
                style_name = (getattr(paragraph.style, "name", "") or "").lower().strip()
            md_prefix = _HEADING_STYLE_TO_MARKDOWN.get(style_name)

            if md_prefix:
                # Chuyển Heading Word → Markdown header (e.g. "# 1. Đăng nhập")
                parts.append(f"{md_prefix} {text}")
                continue

            bold_split = _split_bold_label(paragraph)
            if bold_split:
                label, body = bold_split
                parts.append(f"{_BOLD_HEADING_MARKDOWN} {label}")
                if body:
                    parts.append(body)
                continue

            known_split = _match_known_section_label(text)
            if known_split:
                label, body = known_split
                parts.append(f"{_BOLD_HEADING_MARKDOWN} {label}")
                if body:
                    parts.append(body)
                continue

            parts.append(text)

        for table in document.tables:
            rows = table.rows
            if not rows:
                continue

            header_cells = [cell.text.strip().replace("\n", " ") for cell in rows[0].cells]
            has_header = any(header_cells)
            data_rows = rows[1:] if has_header else rows

            table_lines: list[str] = []
            if has_header:
                table_lines.append(" | ".join(header_cells))

            for row in data_rows:
                cells = [cell.text.strip().replace("\n", " ") for cell in row.cells]
                if not any(cells):
                    continue
                if has_header:
                    # Gắn tên cột vào từng ô ("Cột: giá trị") để mỗi dòng vẫn giữ được
                    # ý nghĩa dù bị tách khỏi hàng header khi text splitter cắt chunk
                    # ở giữa bảng.
                    cells = [
                        f"{header_cells[i]}: {value}" if i < len(header_cells) and header_cells[i] else value
                        for i, value in enumerate(cells)
                    ]
                table_lines.append(" | ".join(cells))

            # Giữ cả bảng là MỘT part duy nhất (thay vì mỗi hàng một part rời rạc như
            # trước) để giảm khả năng bị tách rời nhau khi ghép parts lại bằng "\n\n".
            if table_lines:
                parts.append("\n".join(table_lines))

        return "\n\n".join(parts).strip()
