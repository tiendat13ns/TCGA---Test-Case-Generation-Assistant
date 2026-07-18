from app.services.extractors.base import BaseExtractor

# Mapping từ Word Heading style → Markdown prefix
_HEADING_STYLE_TO_MARKDOWN = {
    "heading 1": "#",
    "heading 2": "##",
    "heading 3": "###",
    "heading 4": "####",
}


class DocxExtractor(BaseExtractor):
    """Trích xuất text từ file .docx, giữ nguyên cấu trúc Heading dưới dạng Markdown.

    Chiến lược:
    - Các đoạn văn mang Word style Heading 1/2/3/4 được chuyển thành
      thẻ Markdown # / ## / ### / #### để MarkdownHeaderTextSplitter có thể
      nhận diện ranh giới chunk sau này.
    - Nội dung bảng (table) vẫn được giữ nguyên dạng pipe-delimited.
    """

    def extract(self, file_path: str) -> str:
        try:
            from docx import Document
        except ImportError as exc:
            raise RuntimeError("python-docx is required to extract DOCX files") from exc

        try:
            document = Document(file_path)
        except Exception as exc:
            raise ValueError(f"Could not read DOCX file: {exc}") from exc

        parts: list[str] = []

        for paragraph in document.paragraphs:
            text = paragraph.text.strip()
            if not text:
                continue

            style_name = ""
            if paragraph.style:
                style_name = (getattr(paragraph.style, "name", "") or "").lower().strip()
            md_prefix = _HEADING_STYLE_TO_MARKDOWN.get(style_name)

            if md_prefix:
                # Chuyển Heading Word → Markdown header (e.g. "# 1. Đăng nhập")
                parts.append(f"{md_prefix} {text}")
            else:
                parts.append(text)

        for table in document.tables:
            for row in table.rows:
                cells = [cell.text.strip().replace("\n", " ") for cell in row.cells]
                if any(cells):
                    parts.append(" | ".join(cells))

        return "\n\n".join(parts).strip()
